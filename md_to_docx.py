import re
import sys
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(section):
    """Sets margins to the Gold Standard: Left 3cm, Right 2.5cm, Top 2.5cm, Bottom 2.0cm."""
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

def set_normal_style(document):
    """Configures the Normal style: Times New Roman, 12pt, Justified, 1.5 spacing."""
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0) # Clean spacing

def set_heading_styles(document):
    """Configures Heading styles."""
    # Heading 1: 16pt Bold, Spacing Before 18pt, After 6pt
    h1 = document.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = None # Auto color (Black)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    
    # Heading 2: 14pt Bold, Spacing Before 12pt, After 6pt
    h2 = document.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = None
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    # Heading 3: 12pt Bold, Spacing Before 12pt, After 6pt
    h3 = document.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = None
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

def parse_markdown(md_text):
    """Simple Markdown parser."""
    lines = md_text.split('\n')
    parsed_lines = []
    
    code_block = False
    
    for line in lines:
        if line.startswith('```'):
            code_block = not code_block
            continue
            
        if code_block:
            parsed_lines.append({'type': 'code', 'content': line})
            continue
            
        # Images: ![Alt](Src)
        elif line.strip().startswith('![') and '](' in line:
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                caption = match.group(1)
                src = match.group(2)
                parsed_lines.append({'type': 'image', 'src': src, 'caption': caption})
            continue

        # Headers
        if line.startswith('# '):
            parsed_lines.append({'type': 'h1', 'content': line[2:].strip()})
        elif line.startswith('## '):
            parsed_lines.append({'type': 'h2', 'content': line[3:].strip()})
        elif line.startswith('### '):
            parsed_lines.append({'type': 'h3', 'content': line[4:].strip()})
        elif line.startswith('#### '):
            parsed_lines.append({'type': 'h4', 'content': line[5:].strip()})
            
        # List items
        elif line.strip().startswith('* ') or line.strip().startswith('- '):
            parsed_lines.append({'type': 'list_bullet', 'content': line.strip()[2:].strip()})
        elif re.match(r'^\d+\.\s', line.strip()):
            content = re.sub(r'^\d+\.\s', '', line.strip())
            parsed_lines.append({'type': 'list_number', 'content': content})
            
        # Tables (Simple detection)
        elif line.strip().startswith('|'):
             parsed_lines.append({'type': 'table_row', 'content': line.strip()})
             
        # Normal text
        elif line.strip():
            parsed_lines.append({'type': 'paragraph', 'content': line.strip()})
        else:
            parsed_lines.append({'type': 'blank', 'content': ''})
            
    return parsed_lines

def apply_formatting(paragraph, text):
    """
    Applies bold/italic formatting to a paragraph.
    Parses **bold** and *italic* and strips other asterisks.
    """
    # Split by bold first: **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            content = part[2:-2]
            content = content.replace('*', '') 
            run = paragraph.add_run(content)
            run.bold = True
        else:
            # Split by italic: *text*
            subparts = re.split(r'(\*.*?\*)', part)
            for subpart in subparts:
                if subpart.startswith('*') and subpart.endswith('*') and len(subpart) > 2:
                    content = subpart[1:-1]
                    run = paragraph.add_run(content)
                    run.italic = True
                else:
                    # Strip any remaining asterisks
                    clean_text = subpart.replace('*', '')
                    if clean_text:
                        paragraph.add_run(clean_text)

def create_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    
    with open(md_path, 'r') as f:
        md_text = f.read()
        
    parsed_lines = parse_markdown(md_text)
    
    document = Document()
    
    # Apply Global Styles
    set_normal_style(document)
    set_heading_styles(document)
    
    # 1. Setup First Section (Front Matter)
    section = document.sections[0]
    set_margins(section)
    
    # --- Title Page ---
    # Add some vertical space
    for _ in range(5):
        document.add_paragraph()
        
    title_p = document.add_paragraph('LLM Self-Improvement: A Systematic Literature Review')
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.style = 'Title'
    
    document.add_paragraph().alignment = WD_ALIGN_PARAGRAPH.CENTER # Spacer
    
    author_p = document.add_paragraph('Jannik Hiller') # Assuming user name or placeholder
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.style = 'Subtitle'
    
    document.add_page_break()

    # --- Table of Contents ---
    document.add_heading('Table of Contents', level=1)
    
    # Insert TOC Field Code
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar)
    
    run.add_text("Right-click to update Table of Contents")
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    document.add_page_break()

    # We need to detect where "1. Introduction" starts to insert a Section Break
    table_buffer = []
    introduction_started = False

    for line in parsed_lines:
        type = line['type']
        
        # Handle content access safely or based on type
        if type == 'image':
            content = line # Pass the whole dict for image
        else:
            content = line['content']
        
        # Detect Introduction to insert Section Break (Next Page)
        if type == 'h1' and ('1. Introduction' in content or 'Introduction' in content) and not introduction_started:
            introduction_started = True
            # Insert Section Break (Next Page) to restart numbering logic
            p = document.add_paragraph()
            p.add_run().add_break(WD_BREAK.SECTION_NEXT_PAGE)
            new_section = document.sections[-1]
            set_margins(new_section)

        # Handle Tables
        if type == 'table_row':
            table_buffer.append(content)
            continue
        else:
            if table_buffer:
                # Process table
                rows = [row for row in table_buffer if not set(row.replace('|', '').strip()) == {'-'}]
                if rows:
                    cols = len(rows[0].split('|')) - 2
                    if cols > 0:
                        # Add Table Caption (Above)
                        document.add_paragraph('Table', style='Caption')
                        
                        table = document.add_table(rows=len(rows), cols=cols)
                        table.style = 'Table Grid'
                        for i, row_str in enumerate(rows):
                            cells = row_str.split('|')[1:-1]
                            for j, cell_text in enumerate(cells):
                                if j < cols:
                                    # FIX: Use apply_formatting instead of direct assignment
                                    # This parses **bold** and *italic* inside cells
                                    cell_p = table.cell(i, j).paragraphs[0]
                                    apply_formatting(cell_p, cell_text.strip())
                                    
                                    # Set font size for table content to 10pt
                                    for run in cell_p.runs:
                                        run.font.size = Pt(10)
                                        run.font.name = 'Times New Roman'
                table_buffer = []
        
        if type == 'h1':
            document.add_heading(content, level=1)
        elif type == 'h2':
            document.add_heading(content, level=2)
        elif type == 'h3':
            document.add_heading(content, level=3)
        elif type == 'h4':
            document.add_heading(content, level=4)
        elif type == 'list_bullet':
            p = document.add_paragraph(style='List Bullet')
            apply_formatting(p, content)
        elif type == 'list_number':
            p = document.add_paragraph(style='List Number')
            apply_formatting(p, content)
        elif type == 'code':
            p = document.add_paragraph(content)
            p.style = 'No Spacing'
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(10)
        elif type == 'image':
            # Insert Image
            try:
                # Add image centered
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(content['src'], width=Inches(6.0)) # Standard width
                
                # Add Caption below
                caption_p = document.add_paragraph(content['caption'])
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_p.style = 'Caption' # Or Normal with manual formatting
                # Ensure caption font is 10pt
                for run in caption_p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
                    
            except Exception as e:
                print(f"Error adding image {content['src']}: {e}")
                
        elif type == 'paragraph':
            p = document.add_paragraph()
            apply_formatting(p, content)
            
    document.save(docx_path)
    print(f"Saved {docx_path}")

if __name__ == "__main__":
    create_docx("final_paper.md", "final_paper.docx")
